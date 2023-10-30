import tkinter as tk
from tkinter import filedialog
from tkinter import simpledialog
from docx import Document

def open_file():
    file_path = filedialog.askopenfilename()
    if file_path:
        with open(file_path, 'r') as file:
            content = file.read()
            text.delete(1.0, tk.END)
            text.insert(tk.END, content)

def save_file():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
    if file_path:
        document = Document()

        # Get text content
        content = text.get(1.0, tk.END)
        document.add_paragraph(content)

        # Save images
        for image in images:
            document.add_picture(image, width=200)

        # Save the document
        document.save(file_path)

def insert_image():
    file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.gif")])
    if file_path:
        images.append(file_path)
        text.insert(tk.END, f"\n[IMAGE:{len(images)}]\n")

# Create main window
root = tk.Tk()
root.title("Enhanced Text Editor")

# Create text widget
text = tk.Text(root, wrap="word")
text.pack(expand=True, fill="both")

# Create menu bar
menu_bar = tk.Menu(root)

# File menu
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Open", command=open_file)
file_menu.add_command(label="Save", command=save_file)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=root.destroy)
menu_bar.add_cascade(label="File", menu=file_menu)

# Edit menu
edit_menu = tk.Menu(menu_bar, tearoff=0)
edit_menu.add_command(label="Insert Image", command=insert_image)
menu_bar.add_cascade(label="Edit", menu=edit_menu)

# Configure menu bar
root.config(menu=menu_bar)

# List to store image paths
images = []

# Run the application
root.mainloop()
